"""
documentGeneration.py
=====================
EquityForm document-generation backend module.

Fills DOCX templates (Operating Agreement and Subscription Agreement) with
deal data fetched from the application state.  Intended to run inside a
FastAPI / Flask service.

Dependencies:
    pip install python-docx fastapi uvicorn

Quick start:
    uvicorn documentGeneration:app --reload --port 8000
"""

from __future__ import annotations

import re
import copy
import io
from datetime import datetime, date
from typing import Any

from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from fastapi import FastAPI, HTTPException
from fastapi.responses import StreamingResponse
from pydantic import BaseModel

# ---------------------------------------------------------------------------
# FastAPI application
# ---------------------------------------------------------------------------

app = FastAPI(title="EquityForm Document Generation API")

# ---------------------------------------------------------------------------
# Pydantic models (mirrors TypeScript store types)
# ---------------------------------------------------------------------------

class DealModel(BaseModel):
    entityName: str | None = None
    formationState: str | None = None
    effectiveDate: str | None = None
    principalAddress: str | None = None
    gpEntityName: str | None = None
    gpEntityState: str | None = None
    gpSignerName: str | None = None
    gpSignerTitle: str | None = None
    registeredAgentName: str | None = None
    registeredAgentAddress: str | None = None
    dealPurpose: str | None = None
    propertyAddress: str | None = None
    propertyCity: str | None = None
    propertyState: str | None = None
    propertyZip: str | None = None
    propertyLegalDescription: str | None = None
    ein: str | None = None

class OfferingModel(BaseModel):
    offeringExemption: str | None = None
    offeringExemptionRule: str | None = None
    solicitationMethod: str | None = None
    minimumInvestment: float | None = None
    closingDate: str | None = None
    preferredReturnEnabled: bool | None = None
    preferredReturnRate: float | None = None
    preferredReturnType: str | None = None
    irrRate: float | None = None
    gpPromote: float | None = None
    lpResidual: float | None = None
    assetManagementFeeDescription: str | None = None
    acquisitionFeeDescription: str | None = None
    dispositionFeeDescription: str | None = None
    consentThreshold: float | None = None
    refinanceThreshold: float | None = None
    amendmentThreshold: float | None = None
    reportPeriod: str | None = None
    reportFrequencyDays: float | None = None
    disputeResolutionMethod: str | None = None
    disputeResolutionVenue: str | None = None

class BankingModel(BaseModel):
    bankName: str | None = None
    accountName: str | None = None
    accountNumber: str | None = None
    routingNumber: str | None = None

class InvestorModel(BaseModel):
    id: str
    fullLegalName: str | None = None
    subscriberType: str | None = "individual"
    entityLegalName: str | None = None
    entityType: str | None = None
    formationState: str | None = None
    taxId: str | None = None
    streetAddress: str | None = None
    city: str | None = None
    state: str | None = None
    zip: str | None = None
    email: str | None = None
    phone: str | None = None
    subscriptionAmount: float | None = None
    classAUnits: float | None = None
    ownershipPct: float | None = None
    accreditedInvestor: bool | None = None
    signerName: str | None = None
    signerTitle: str | None = None

class AppDataModel(BaseModel):
    deal: DealModel = DealModel()
    offering: OfferingModel = OfferingModel()
    banking: BankingModel = BankingModel()
    investors: list[InvestorModel] = []

class GenerateRequest(BaseModel):
    appData: AppDataModel
    templatePath: str | None = None   # server-side path to .docx template

# ---------------------------------------------------------------------------
# Formatting helpers
# ---------------------------------------------------------------------------

def fmt_date(val: str | date | None, fmt: str = "%B %d, %Y") -> str:
    """Return a human-readable date string, or empty string on failure."""
    if not val:
        return ""
    if isinstance(val, (date, datetime)):
        return val.strftime(fmt)
    # Try ISO format first
    for pattern in ("%Y-%m-%d", "%m/%d/%Y", "%Y-%m-%dT%H:%M:%SZ", "%Y-%m-%dT%H:%M:%S"):
        try:
            return datetime.strptime(str(val)[:19], pattern).strftime(fmt)
        except ValueError:
            pass
    return str(val)

def fmt_currency(val: Any) -> str:
    """Return a formatted dollar amount, e.g. $1,000,000.00"""
    if val is None or val == "":
        return ""
    try:
        return f"${float(val):,.2f}"
    except (ValueError, TypeError):
        return str(val)

def fmt_pct(val: Any) -> str:
    """Return a formatted percentage, e.g. 8%"""
    if val is None or val == "":
        return ""
    try:
        return f"{float(val):.4g}%"
    except (ValueError, TypeError):
        return str(val)

# ---------------------------------------------------------------------------
# Placeholder → field mapping
# ---------------------------------------------------------------------------

def build_oa_values(app_data: AppDataModel) -> dict[str, str]:
    """
    Build a flat dict mapping every OA placeholder token to its resolved value.

    Group A — Entity
    Group B — GP
    Group C — Property
    Group D — Offering economics
    Group E — Operations
    Group F — Banking
    """
    d = app_data.deal
    o = app_data.offering
    b = app_data.banking

    property_full = ", ".join(filter(None, [
        d.propertyAddress,
        d.propertyCity,
        (f"{d.propertyState} {d.propertyZip}".strip() if (d.propertyState or d.propertyZip) else None),
    ]))

    pref_return_text = ""
    if o.preferredReturnEnabled and o.preferredReturnRate:
        pref_type = (o.preferredReturnType or "cumulative").replace("-", " ")
        pref_return_text = f"{fmt_pct(o.preferredReturnRate)} per annum ({pref_type})"

    return {
        # Group A — Entity
        "ENTITY_NAME":              d.entityName or "",
        "FORMATION_STATE":          d.formationState or "",
        "EFFECTIVE_DATE":           fmt_date(d.effectiveDate),
        "PRINCIPAL_ADDRESS":        d.principalAddress or "",
        "EIN":                      d.ein or "",
        "REGISTERED_AGENT_NAME":    d.registeredAgentName or "",
        "REGISTERED_AGENT_ADDRESS": d.registeredAgentAddress or "",
        # Group B — GP
        "GP_ENTITY_NAME":           d.gpEntityName or "",
        "GP_ENTITY_STATE":          d.gpEntityState or "",
        "GP_SIGNER_NAME":           d.gpSignerName or "",
        "GP_SIGNER_TITLE":          d.gpSignerTitle or "",
        # Group C — Property
        "PROPERTY_ADDRESS":         d.propertyAddress or "",
        "PROPERTY_CITY":            d.propertyCity or "",
        "PROPERTY_STATE":           d.propertyState or "",
        "PROPERTY_ZIP":             d.propertyZip or "",
        "PROPERTY_FULL":            property_full,
        "PROPERTY_LEGAL_DESC":      d.propertyLegalDescription or "",
        "DEAL_PURPOSE":             d.dealPurpose or "",
        # Group D — Economics
        "OFFERING_EXEMPTION":       o.offeringExemption or "",
        "OFFERING_EXEMPTION_RULE":  o.offeringExemptionRule or "",
        "SOLICITATION_METHOD":      o.solicitationMethod or "",
        "MIN_INVESTMENT":           fmt_currency(o.minimumInvestment),
        "CLOSING_DATE":             fmt_date(o.closingDate),
        "PREFERRED_RETURN":         pref_return_text,
        "PREFERRED_RETURN_RATE":    fmt_pct(o.preferredReturnRate),
        "PREFERRED_RETURN_TYPE":    (o.preferredReturnType or "").replace("-", " "),
        "IRR_RATE":                 fmt_pct(o.irrRate),
        "GP_PROMOTE":               fmt_pct(o.gpPromote),
        "LP_RESIDUAL":              fmt_pct(o.lpResidual),
        # Group E — Operations
        "ASSET_MGMT_FEE":           o.assetManagementFeeDescription or "",
        "ACQUISITION_FEE":          o.acquisitionFeeDescription or "",
        "DISPOSITION_FEE":          o.dispositionFeeDescription or "",
        "CONSENT_THRESHOLD":        fmt_pct(o.consentThreshold),
        "REFINANCE_THRESHOLD":      fmt_pct(o.refinanceThreshold),
        "AMENDMENT_THRESHOLD":      fmt_pct(o.amendmentThreshold),
        "REPORT_PERIOD":            o.reportPeriod or "",
        "REPORT_FREQUENCY_DAYS":    str(int(o.reportFrequencyDays)) if o.reportFrequencyDays else "",
        "DISPUTE_RESOLUTION":       o.disputeResolutionMethod or "",
        "DISPUTE_VENUE":            o.disputeResolutionVenue or "",
        # Group F — Banking
        "BANK_NAME":                b.bankName or "",
        "BANK_ACCOUNT_NAME":        b.accountName or "",
        "BANK_ACCOUNT_NUMBER":      b.accountNumber or "",
        "BANK_ROUTING_NUMBER":      b.routingNumber or "",
    }

def build_sub_values(investor: InvestorModel, app_data: AppDataModel, oa_values: dict[str, str]) -> dict[str, str]:
    """
    Build the per-investor placeholder dict for a Subscription Agreement.
    Inherits all OA values and overrides with investor-specific fields.
    """
    total_units = sum(i.classAUnits or 0 for i in app_data.investors)
    ownership_pct = (
        ((investor.classAUnits or 0) / total_units * 100) if total_units > 0 else 0
    )

    investor_address = ", ".join(filter(None, [
        investor.streetAddress,
        investor.city,
        (f"{investor.state} {investor.zip}".strip() if (investor.state or investor.zip) else None),
    ]))

    signer_name  = investor.signerName  or investor.fullLegalName or ""
    signer_title = investor.signerTitle or ("Managing Member" if investor.subscriberType == "entity" else "")

    vals = {**oa_values}
    vals.update({
        "INVESTOR_NAME":          investor.fullLegalName or "",
        "INVESTOR_TYPE":          investor.subscriberType or "individual",
        "INVESTOR_ENTITY_NAME":   investor.entityLegalName or "",
        "INVESTOR_ENTITY_TYPE":   investor.entityType or "",
        "INVESTOR_FORMATION_STATE": investor.formationState or "",
        "INVESTOR_TAX_ID":        investor.taxId or "",
        "INVESTOR_ADDRESS":       investor_address,
        "INVESTOR_EMAIL":         investor.email or "",
        "INVESTOR_PHONE":         investor.phone or "",
        "INVESTOR_SIGNER_NAME":   signer_name,
        "INVESTOR_SIGNER_TITLE":  signer_title,
        "SUBSCRIPTION_AMOUNT":    fmt_currency(investor.subscriptionAmount),
        "CLASS_A_UNITS":          str(int(investor.classAUnits)) if investor.classAUnits else "0",
        "OWNERSHIP_PCT":          fmt_pct(round(ownership_pct, 4)),
        "ACCREDITED_STATUS":      "accredited investor" if investor.accreditedInvestor else "non-accredited investor",
    })
    return vals

# ---------------------------------------------------------------------------
# Core template fill functions
# ---------------------------------------------------------------------------

def _replace_in_paragraph(paragraph: Any, values: dict[str, str]) -> None:
    """
    Replace all [TOKEN] placeholders in a paragraph, handling run-splitting.

    python-docx can split a token across multiple runs (e.g. "[EN" in run 1,
    "TITY_NAME]" in run 2).  We join the paragraph text, find tokens, then
    rebuild the run list.
    """
    full_text = "".join(run.text for run in paragraph.runs)
    if "[" not in full_text:
        return

    def replacer(match: re.Match) -> str:
        token = match.group(1)
        return values.get(token, match.group(0))  # leave unknown tokens as-is

    new_text = re.sub(r"\[([A-Z0-9_]+)\]", replacer, full_text)

    if new_text == full_text:
        return

    # Preserve formatting from the first run, clear the rest
    if paragraph.runs:
        paragraph.runs[0].text = new_text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(new_text)


def fill_template(template_path: str | None, values: dict[str, str]) -> Document:
    """
    Open *template_path* (or create a minimal stub if None/missing) and fill
    every [TOKEN] placeholder in paragraphs, table cells, headers and footers.
    Returns the modified Document object.
    """
    try:
        doc = Document(template_path) if template_path else Document()
    except Exception:
        doc = Document()

    def _process(doc_obj: Document) -> None:
        # Body paragraphs
        for para in doc_obj.paragraphs:
            _replace_in_paragraph(para, values)
        # Tables
        for table in doc_obj.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        _replace_in_paragraph(para, values)
        # Headers and footers
        for section in doc_obj.sections:
            for part in (section.header, section.footer):
                if part:
                    for para in part.paragraphs:
                        _replace_in_paragraph(para, values)

    _process(doc)
    return doc


def remove_dev_notes(doc: Document) -> Document:
    """
    Remove any paragraph that starts with "DEV NOTE:" or "[DEV]" from the document.
    These are template author annotations not intended for the final document.
    """
    to_remove = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text.startswith("DEV NOTE:") or text.startswith("[DEV]"):
            to_remove.append(para)

    for para in to_remove:
        p = para._element
        p.getparent().remove(p)

    return doc


def doc_to_bytes(doc: Document) -> bytes:
    """Serialize a Document to bytes for HTTP streaming."""
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()

# ---------------------------------------------------------------------------
# FastAPI endpoints
# ---------------------------------------------------------------------------

@app.post(
    "/api/deals/{deal_id}/generate/oa",
    response_class=StreamingResponse,
    summary="Generate Operating Agreement DOCX",
)
async def generate_oa(deal_id: str, req: GenerateRequest) -> StreamingResponse:
    """Fill the OA template and return a DOCX binary."""
    oa_values = build_oa_values(req.appData)
    doc       = fill_template(req.templatePath, oa_values)
    doc       = remove_dev_notes(doc)
    content   = doc_to_bytes(doc)
    filename  = f"OperatingAgreement_{deal_id}.docx"
    return StreamingResponse(
        io.BytesIO(content),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


@app.post(
    "/api/deals/{deal_id}/generate/sub/{investor_id}",
    response_class=StreamingResponse,
    summary="Generate Subscription Agreement DOCX for one investor",
)
async def generate_sub(deal_id: str, investor_id: str, req: GenerateRequest) -> StreamingResponse:
    """Fill the Sub Agreement template for a specific investor and return DOCX."""
    investor = next(
        (i for i in req.appData.investors if i.id == investor_id),
        None,
    )
    if not investor:
        raise HTTPException(status_code=404, detail=f"Investor {investor_id} not found")

    oa_values  = build_oa_values(req.appData)
    sub_values = build_sub_values(investor, req.appData, oa_values)
    doc        = fill_template(req.templatePath, sub_values)
    doc        = remove_dev_notes(doc)
    content    = doc_to_bytes(doc)
    name_slug  = re.sub(r"[^a-zA-Z0-9_-]", "_", investor.fullLegalName or investor_id)
    filename   = f"SubscriptionAgreement_{deal_id}_{name_slug}.docx"
    return StreamingResponse(
        io.BytesIO(content),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )
